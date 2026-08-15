import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_title_page():
    out_dir = "/Users/shamthakare/.gemini/antigravity/scratch/quorumshift/submission/ieee_tai/resubmission"
    os.makedirs(out_dir, exist_ok=True)
    out_docx = os.path.join(out_dir, "AdaptiveReplica_TAI_Title_Page.docx")

    doc = Document()
    
    # Set 1-inch margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Manuscript ID Header
    p_meta = doc.add_paragraph()
    r_meta = p_meta.add_run("IEEE TRANSACTIONS ON ARTIFICIAL INTELLIGENCE — TITLE PAGE\nManuscript ID: TAI-2026-Aug-A-01875")
    r_meta.font.name = "Calibri"
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = RGBColor(100, 116, 139)
    p_meta.paragraph_format.space_after = Pt(16)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("AdaptiveReplica: Dynamic Quorum Adaptation and Failure-Aware Replica Selection in Distributed Consensus")
    r_title.bold = True
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(16)

    # Author
    p_auth = doc.add_paragraph()
    r_auth_h = p_auth.add_run("Author:\n")
    r_auth_h.bold = True
    r_auth_h.font.size = Pt(12)
    r_auth = p_auth.add_run("Sham Satish Thakare\nIndependent Researcher\nPune, Maharashtra, India")
    r_auth.font.size = Pt(11)
    p_auth.paragraph_format.space_after = Pt(14)

    # Correspondence
    p_corr = doc.add_paragraph()
    r_corr_h = p_corr.add_run("Corresponding Author & Contact Information:\n")
    r_corr_h.bold = True
    r_corr_h.font.size = Pt(12)
    r_corr = p_corr.add_run(
        "Sham Satish Thakare\n"
        "Independent Researcher\n"
        "Email: shamthakare3000@gmail.com\n"
        "Website / Code Repository: https://github.com/shamddd/quorumshift"
    )
    r_corr.font.size = Pt(11)
    p_corr.paragraph_format.space_after = Pt(14)

    # Acknowledgments & IEEE AI Disclosure
    p_ack = doc.add_paragraph()
    r_ack_h = p_ack.add_run("Acknowledgments & IEEE AI Disclosure Statement:\n")
    r_ack_h.bold = True
    r_ack_h.font.size = Pt(12)
    r_ack = p_ack.add_run(
        "In accordance with IEEE Author Guidelines and IEEE Publication Services and Products Board (PSPB) Operations Manual policies regarding generative AI systems: "
        "Generative AI assistance (Google Antigravity / Gemini models) was utilized during code refactoring, simulation monitoring, and LaTeX template formatting. "
        "All distributed consensus algorithms, C++ state machine implementations, fault injection experiments, and formal invariant proofs were independently designed, executed, verified, and audited by the author. "
        "The author assumes full responsibility for the integrity, validity, and accuracy of all scientific claims and results presented in this manuscript.\n\n"
        "Funding Declaration: This research received no external grant funding.\n"
        "Conflict of Interest: The author declares no competing financial or non-financial interests."
    )
    r_ack.font.size = Pt(10.5)

    doc.save(out_docx)
    print(f"Saved {out_docx}")

if __name__ == "__main__":
    create_title_page()
